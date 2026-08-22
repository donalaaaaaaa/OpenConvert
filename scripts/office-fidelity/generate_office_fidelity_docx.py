"""Generate the deterministic DOCX fixture used by Office fidelity tests."""

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Noto Sans CJK SC"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
GRID = "A9B7C6"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths):
                cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=9)


def build(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 12, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("OpenConvert 保真度回归 · 中文页眉 DOCX-HEADER"), 9, True, "5B6573")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("中文页脚 DOCX-FOOTER · 第 "), 9, False, "5B6573")
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), 9, False, "5B6573")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    set_run_font(title.add_run("OpenConvert Office 中文保真度样本"), 22, True, "0B2545")

    intro = doc.add_paragraph()
    set_run_font(
        intro.add_run(
            "DOCX-中文-甲：离线转换应保留中文标点、数字 2026、英文 OpenConvert，"
            "并正确处理中文字体回退。"
        ),
        11,
    )

    doc.add_heading("复杂表格", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    merged = table.cell(0, 0).merge(table.cell(0, 3))
    merged.text = "DOCX-表格-合计 · 区域营收明细"
    shade(merged, LIGHT_BLUE)
    for run in merged.paragraphs[0].runs:
        set_run_font(run, 11, True, "0B2545")
    headers = ["区域", "一季度", "二季度", "全年目标"]
    for col, value in enumerate(headers):
        cell = table.cell(1, col)
        cell.text = value
        shade(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 10, True, "FFFFFF")

    rows = [
        ("华东", "¥1,280,000", "¥1,460,000", "¥6,000,000"),
        ("华南", "¥980,000", "¥1,120,000", "¥4,800,000"),
        ("西部", "¥760,000", "¥830,000", "¥3,600,000"),
        ("合计", "¥3,020,000", "¥3,410,000", "¥14,400,000"),
    ]
    for row_values in rows:
        cells = table.add_row().cells
        for col, value in enumerate(row_values):
            cells[col].text = value
            for run in cells[col].paragraphs[0].runs:
                set_run_font(run, 10, col == 0)
    set_table_geometry(table, [1800, 2520, 2520, 2520])

    doc.add_page_break()
    doc.add_heading("跨页内容与页眉页脚", level=1)
    paragraph = doc.add_paragraph()
    set_run_font(
        paragraph.add_run(
            "DOCX-第二页-乙：本页用于确认分页后中文页眉与页脚仍存在，"
            "正文字符不应出现方框、缺字或异常换行。"
        ),
        11,
    )
    doc.add_heading("验收清单", level=2)
    for text in (
        "中文字体可读，关键文本 DOCX-CJK-PASS 可提取。",
        "表格标题、表头和货币数字保持完整。",
        "页眉 DOCX-HEADER 与页脚 DOCX-FOOTER 出现在导出 PDF。",
    ):
        item = doc.add_paragraph(style="List Bullet")
        set_run_font(item.add_run(text), 11)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: generate_office_fidelity_docx.py OUTPUT.docx")
    build(Path(sys.argv[1]).resolve())
